from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "docs/manual/BatteryCurrent_User_Manual_Preliminary.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E2EC", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("BatteryCurrent User Manual")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Preliminary guide for test users")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("555555")

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    run = note.add_run("Draft status: ")
    run.bold = True
    note.add_run(
        "This manual is intended as a first-pass guide. Screenshot placeholders are included for later replacement or annotation."
    )


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_placeholder(doc, title, instructions):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    set_table_width(table, 9360)
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F2F4F7")
    set_cell_border(cell, "C7D2E1", "12")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(f"[ Screenshot placeholder: {title} ]")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")
    paragraph.add_run("\n")
    detail = paragraph.add_run(instructions)
    detail.font.size = Pt(10)
    detail.font.color.rgb = RGBColor.from_string("555555")
    doc.add_paragraph()


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    set_table_width(table, 9360)
    set_cell_width(table.rows[0].cells[0], 2700)
    set_cell_width(table.rows[0].cells[1], 6660)

    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "What it does"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True

    for key, value in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], 2700)
        set_cell_width(cells[1], 6660)
        cells[0].text = key
        cells[1].text = value
        for cell in cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_footer(doc):
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("BatteryCurrent preliminary manual")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("777777")


def build_manual():
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "BatteryCurrent is an Android overlay utility for monitoring live battery behavior while using the phone normally. "
        "It shows charging or discharging current, battery temperature, voltage, accumulated charge or energy, battery percentage, and a live history graph."
    )
    add_bullets(doc, [
        "The foreground display is intentionally compact so it can sit over other apps.",
        "The graph popup provides detailed history, zooming, selectable secondary traces, and control buttons.",
        "Sampling is designed to be low power, updating about every 10 seconds.",
        "Battery capacity event logging is handled in the background when qualifying charge or discharge sessions occur.",
    ])

    doc.add_heading("2. Starting the App", level=1)
    doc.add_paragraph(
        "When BatteryCurrent starts, the startup page provides monitoring controls, foreground display options, original battery capacity entry, "
        "and a calibration setup launcher. Starting monitoring opens the foreground display and shows the current live values. "
        "On the first run the foreground display opens near the center of the display. On later runs it reopens at the last dragged position."
    )
    add_numbered(doc, [
        "Open BatteryCurrent from the app launcher.",
        "Grant the overlay permission if Android asks for it.",
        "Use Start monitoring to show the foreground display, or use Calibration setup when performing an occasional controlled calibration measurement.",
        "Tap the foreground display to open the graph popup.",
    ])
    add_key_value_table(doc, [
        ("X button", "Closes the startup page without changing the monitoring state."),
        ("Reset foreground display", "Moves the floating display back to the center if it was dragged to an unreachable edge."),
        ("Calibration setup", "Starts the guided calibration flow. The button reads Start when inactive and ON while armed or running."),
        ("Light-theme foreground display", "Switches the foreground overlay and chart popups to a light beige theme with higher-contrast graph, table, button, and text colors."),
        ("Reset graph at capacity window", "Automatically starts a fresh graph segment when charging crosses up through the low threshold or discharging crosses down through the high threshold. This reset affects the graph display baseline, not the stored capacity history."),
        ("Capacity window", "Sets the low and high battery percentages used for normal capacity estimates, status dots, graph reference lines, and optional graph reset points. The low threshold must be at least 20%, and the window must be at least 40% wide. For example, a 40% to 80% window uses measured mAh multiplied by 2.5."),
        ("Original capacity", "Stores the phone's rated battery capacity so the app can color capacity estimates relative to the original rating."),
        ("Monitor", "Starts or stops the normal monitoring service."),
    ])

    add_placeholder(
        doc,
        "Startup screen",
        "Insert a screenshot showing the startup page with Start monitoring, Calibration setup, foreground display reset, original capacity, and close X.",
    )

    add_placeholder(
        doc,
        "Foreground display",
        "Insert a screenshot showing the small floating overlay with time, current, temperature, voltage, energy or charge, and battery percent.",
    )

    doc.add_heading("3. Foreground Display", level=1)
    doc.add_paragraph(
        "The foreground display is the small live readout shown over other apps. It can be dragged to a convenient location unless it is locked."
    )
    add_key_value_table(doc, [
        ("Time", "Elapsed time since the current measurement session began or was reset."),
        ("Current", "Live averaged battery current in mA. Positive and negative values indicate direction of battery flow."),
        ("Temperature", "Battery temperature in degrees C or degrees F."),
        ("Voltage", "Measured battery voltage from Android battery status data."),
        ("Energy / Charge", "Accumulated mWh or mAh since reset, depending on the selected unit."),
        ("Battery", "Battery state of charge as reported by Android. Battery percent is green above 30%, amber from 15% to 30%, and red below 15%."),
        ("CAL prefix", "Appears during calibration so the user knows the stricter measurement mode is active."),
        ("Status dot", "A yellow dot means a capacity event is armed. A blinking green dot means a normal capacity event is recording. A blinking red dot means calibration is active."),
    ])

    doc.add_heading("4. Graph Popup", level=1)
    doc.add_paragraph(
        "Tap the foreground display to open the graph popup. The popup shows the main accumulated mAh or mWh trace and an optional right-axis trace."
    )
    doc.add_paragraph(
        "If Light-theme foreground display is enabled on the startup page, the graph popup, SOC curve popup, charge-history table, event details, and calibration summary use a light chart theme. "
        "The light theme keeps stronger text, button, axis, and trace colors so the chart remains readable on the beige background."
    )
    add_key_value_table(doc, [
        ("x button", "Closes the graph popup and returns to the foreground display."),
        ("Background", "Hides the foreground overlay while keeping the service running in the notification shade."),
        ("Stop", "Stops monitoring. Use this when you want to fully stop the app."),
        ("Lock / Unlock", "Locks or unlocks the foreground display position."),
        ("mAh / mWh", "Switches the main graph and accumulated readout between charge and energy."),
        ("12h / 24h", "Switches clock-time x-axis labels between 12-hour and 24-hour format. Long-press the x-axis to switch between elapsed time and clock time."),
        ("Clr Data", "Opens a confirmation prompt to reset the accumulated mAh or mWh value to zero and start a new graph timeline."),
        ("Stats", "Opens a separate statistics page with equivalent charge/discharge cycles, average C-rate, raw capacity, capacity normalized to a 0.2C reference current, and shortcuts to Charge History and SOC Linearity."),
        ("degrees button", "Switches temperature display between degrees C and degrees F."),
        ("Display toggles", "Turn individual foreground-display fields on or off. The notification keeps showing the full live value set."),
        ("Collapse arrow", "Collapses or expands the menu buttons so the user can leave only the graph visible."),
    ])

    add_placeholder(
        doc,
        "Graph popup",
        "Insert a screenshot showing the graph popup, controls, right-axis selector, FD/adjusted capacity line, and Reset Zoom button if visible.",
    )

    add_placeholder(
        doc,
        "Collapsed graph menu",
        "Insert a screenshot showing the graph popup with the button menu collapsed so only the graph and compact summary remain visible.",
    )

    doc.add_heading("5. Reading the Graph", level=1)
    doc.add_paragraph(
        "The left axis shows accumulated mAh or mWh. Rising green sections indicate increasing accumulated value; falling red sections indicate decreasing accumulated value. "
        "The right axis is selectable and can show battery percent, temperature, voltage, or current. "
        "The graph can also be used to verify charging profiles, such as smart charging behavior and manufacturer charging claims."
    )
    add_bullets(doc, [
        "Tap the grey right-axis label to cycle through Batt %, Temp, Volt, and mA.",
        "Battery percent, voltage, temperature, and current auto-scale for the visible data. Battery percent can extend above 100% only when manually zoomed.",
        "For mA mode, positive current is light green and negative current is orange, with a dashed zero-current reference line.",
        "Current, voltage, and temperature traces include a thin, lighter cumulative running-average trace. It begins with the first graph sample and converges toward the true average as more samples are collected.",
        "When zoomed out, voltage, temperature, and current traces are smoothed for readability without changing the stored data.",
        "When Batt % is selected, faint reference lines mark the configured low and high capacity-event thresholds.",
        "Long-press the x-axis labels to switch between elapsed time and actual clock time. The 12h/24h button controls clock label format.",
    ])

    doc.add_heading("6. Zooming and Moving the Graph", level=1)
    doc.add_paragraph(
        "The popup itself and graph zooming use deliberate controls so the user can move the popup without accidentally changing the graph scale."
    )
    add_key_value_table(doc, [
        ("One-finger drag", "Moves the popup window."),
        ("X-axis zoom button", "Selects horizontal time zoom. The selected zoom axis is highlighted."),
        ("Y-axis zoom button", "Selects vertical/right-axis zoom. Use this when the right-axis trace needs more detail."),
        ("Two-finger pinch", "Zooms the selected axis. Pinch gestures can be performed over the popup area for easier control."),
        ("Two-finger drag", "Pans the selected zoomed axis while the popup itself remains fixed during zoom mode."),
        ("Reset Zoom", "Appears after zooming or panning. Tap it to return both axes to the normal auto view."),
    ])

    doc.add_heading("7. Background Operation and Notification", level=1)
    doc.add_paragraph(
        "BatteryCurrent can continue running in the background. In this state, the foreground overlay is hidden, but the notification remains available from the Android pull-down shade."
    )
    add_bullets(doc, [
        "Use Background from the graph popup to hide the overlay while monitoring continues.",
        "Use the notification to bring the foreground display back.",
        "Use Stop from the graph popup when you want to end monitoring completely.",
    ])

    doc.add_heading("8. Battery Capacity Event Logging", level=1)
    doc.add_paragraph(
        "BatteryCurrent records qualifying charge and discharge sessions in the app private folder. Normal capacity estimates use completed sessions that span the user-selected capacity window. "
        "The default window is 25% to 75%, but it can be changed on the startup page. "
        "These estimates are stored and used as trend data, but they play a secondary role once a calibration result exists."
    )
    add_bullets(doc, [
        "A charge event starts near the low threshold and completes when it reaches the high threshold.",
        "A discharge event starts near the high threshold and completes when it reaches the low threshold.",
        "Plugging or unplugging during a qualifying session cancels that session and waits for the next threshold crossing.",
        "Daily capacity estimates are stored separately and updated only from completed qualifying events.",
        "When a calibration result exists, the chart capacity line shows both the fixed calibration value and an adjusted value that reflects later trend movement.",
        "Tap the Stats button to review equivalent full cycles and capacity statistics. The 0.2C-normalized capacity uses the battery capacity reference and learned load sensitivity so low-current discharge is not treated as directly equivalent to higher-current discharge.",
        "The Stats page also contains shortcuts for Charge History and SOC Linearity, keeping the main graph controls less crowded.",
        "If no calibration result exists yet, the chart falls back to showing the configured-window extrapolated estimate.",
    ])

    doc.add_heading("9. Calibration Setup", level=1)
    doc.add_paragraph(
        "Calibration is the most accurate battery-capacity measurement mode. It is intended for occasional controlled testing, not daily use, because deep discharge cycles add battery stress."
    )
    add_numbered(doc, [
        "Charge the phone until Android reports 100%.",
        "Leave the phone connected long enough to fully top off.",
        "Open BatteryCurrent and tap Calibration setup, then press Start while Android still reports 100%.",
        "Disconnect the charger after Start.",
        "The startup page closes and the graph opens while calibration is armed.",
        "The app resets the graph/mAh reading and starts measuring automatically when the phone reaches 99%.",
        "Leave monitoring running until the phone discharges to 15%. The app then records the completed calibration.",
    ])
    add_bullets(doc, [
        "Calibration measures the 99% to 15% discharge range and computes capacity as discharged mAh divided by 0.84.",
        "If the charger is connected after measurement begins, the service is stopped, or the measurement is interrupted, the calibration is stopped and no incomplete row is written.",
        "Completed calibration rows are stored in battery_calibration_tests.csv with start and end times, discharged mAh, capacity estimate, average temperature, voltage, and current.",
        "The chart shows Calibration battery capacity [date]: raw mAh and Adj: adjusted mAh. When the rated/original capacity is entered, the adjusted value also shows the percent difference from that rating. The raw value is the fixed result from the last controlled calibration. The adjusted value is a trend-aware estimate that moves gradually as later everyday-use measurements indicate capacity drifting up or down.",
        "Tap the calibration capacity line in the chart to open a compact table of the latest calibration result.",
    ])

    add_placeholder(
        doc,
        "Calibration setup dialog",
        "Insert a screenshot showing the calibration instructions with the 100% requirement, disconnect reminder, Start button, and Cancel button.",
    )

    doc.add_heading("10. SOC Curve and Load Sensitivity", level=1)
    doc.add_paragraph(
        "The SOC Curve view helps diagnose how linear the phone's reported battery percentage is compared with measured mAh use."
    )
    add_bullets(doc, [
        "SOC bucket data is collected during normal discharge, including during calibration.",
        "The SOC curve plots individual bucket samples as dots and draws a piecewise linear line through the average for each bucket.",
        "Outlier points can reveal noisy battery-percentage reporting, load effects, or areas where the phone's fuel gauge is less linear.",
        "The load sensitivity line reports k with two decimals and a short comment such as k=1.00 (excellent), k=1.12 (mild), k=1.22 (noticeable), k=1.35 (high sag), or k>1.50 (check data).",
        "This load-sensitivity value is best treated as a phone-specific load sensitivity index, not a laboratory Peukert constant.",
    ])

    add_placeholder(
        doc,
        "SOC curve popup",
        "Insert a screenshot showing SOC bucket dots, the averaged piecewise line, and the deviation axis.",
    )

    doc.add_heading("11. Data and Privacy", level=1)
    doc.add_paragraph(
        "BatteryCurrent stores its measurement history and battery capacity event files in the app private data area. "
        "The app is intended to monitor battery behavior locally. It does not need cloud storage for normal operation."
    )
    add_bullets(doc, [
        "Graph history is capped to reduce storage growth.",
        "Capacity event, calibration, SOC bucket, and moving-average files are stored as CSV files for later inspection.",
        "Private app data should survive normal in-place app updates, but uninstalling the app or installing with a mismatched signing key can remove the stored data.",
        "Deleting app data or uninstalling the app removes the stored history.",
    ])

    doc.add_heading("12. Troubleshooting", level=1)
    add_key_value_table(doc, [
        ("Overlay does not appear", "Check Android overlay permission for BatteryCurrent, then reopen the app."),
        ("Foreground display cannot move", "Use Reset foreground display to centre from the startup page, or open the graph popup and tap Unlock if the overlay is locked."),
        ("Graph looks crowded", "Use the X or Y zoom controls, pinch to zoom the selected axis, or collapse the menu buttons."),
        ("No capacity estimate yet", "The app needs completed qualifying charge or discharge sessions before daily estimates are available."),
        ("No calibration result yet", "Run the guided calibration from 99% down to 15%. Until then, the chart uses the configured-window estimate when available."),
        ("Calibration will not start", "Start only arms calibration when Android reports 100%. Charge fully, press Start while the battery still reads 100%, then disconnect the charger. The app begins measuring automatically at 99%."),
        ("Monitor button looks stale", "The startup page checks the live service heartbeat. If Android killed the service, the button returns to Start/OFF when reopened."),
        ("App was stopped", "Restart BatteryCurrent from the launcher. If the service was fully stopped, it cannot measure the period while it was not running."),
    ])

    doc.add_heading("13. Notes for This Draft", level=1)
    doc.add_paragraph(
        "This manual is preliminary. Before public release, update screenshots, confirm terminology, add Play Store installation notes, and verify any privacy wording against the final app behavior."
    )

    add_footer(doc)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_manual()
